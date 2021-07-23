import os
from datetime import datetime
import re

from docx import Document
import pandas as pd
import extract_msg
from tqdm import tqdm

import config


# Find the total number of files before parsing for keywords
def find_number_of_files():
    file_count = 0
    for file_type in config.file_types:
        configfiles = [os.path.join(dirpath, f)
                       for dirpath, dirnames, files in os.walk(config.path)
                       for f in files if f.endswith(file_type)]

        for file in configfiles:
            file_count += 1
    return file_count


# Created as a check. A documents date should not change during the parsing. If it does, there is an error
def check_dates_match(document_name, date_created, date_modified):
    if len(document_name) > 1 and document_name[-1] == document_name[-2]:
        if date_created[-1] != date_created[-2] or date_modified[-1] != date_modified[-2]:
            print("Something has gone wrong, the dates for the same document are not the same.")
            sys.exit()


# Get metadata from the .docx document
def getMetaData(document):
    metadata = {}
    prop = document.core_properties
    metadata["author"] = prop.author
    metadata["created"] = prop.created
    metadata["identifier"] = prop.identifier
    metadata["modified"] = prop.modified
    return metadata


# Checks if the keyword is in the docx title
def document_title_contains_kw(word_to_search, file):
    if word_to_search in file.lower():
        return True
    else:
        return False


# Adds 1 to the word count if word is in document/email
def add_to_word_count(words, word_to_search):
    for word in words:
        if word_to_search in word.lower():
            return 1


# The main function, searches document/email for keywords
def keyword_search(file_count):
    # Sets up range of blank lists ready to be appended to
    document_name = []
    document_name_contains_word = []
    count_of_word_in_document = []
    type_of_file = []
    word_searched = []
    folder_searched = []
    date_created = []
    date_modified = []
    author = []
    errors = []
    pbar = tqdm(total=file_count)

    for file_type in config.file_types:
        # Finds all files including in subfolders
        configfiles = [os.path.join(dirpath, f)
                       for dirpath, dirnames, files in os.walk(config.path)
                       for f in files if f.endswith(file_type)]

        for file in configfiles:
            pbar.update(1)
            try:
                if file_type == '.docx':
                    document = Document(file)
                    metadata_dict = getMetaData(document)

                    for word_to_search in config.words_to_search:
                        type_of_file.append(file_type)
                        folder_searched.append(os.path.basename(config.path))
                        date_created.append(metadata_dict.get('created').strftime('%d/%m/%Y'))
                        if metadata_dict.get('modified') != None:
                            date_modified.append(metadata_dict.get('modified').strftime('%d/%m/%Y'))
                        else:
                            date_modified.append('')
                        author.append(metadata_dict.get('author'))
                        document_name.append(file.replace(config.path, ''))
                        word_searched.append(word_to_search)

                        check_dates_match(document_name, date_created, date_modified)
                        count_of_word_in_text_tmp = 0
                        count_of_word_in_tables_tmp = 0
                        count_of_word_in_headers_and_footers_tmp = 0
                        document_name_contains_word.append(document_title_contains_kw(word_to_search, file))

                        # docx files are in paragraphs so we have to parse each paragraph.
                        if config.search_text == True:
                            for paragraph in document.paragraphs:
                                words = paragraph.text.split()
                                for word in words:
                                    if word_to_search == word.lower():
                                        count_of_word_in_text_tmp += 1

                        # tables have to be parsed differently.
                        if config.search_tables == True:
                            for table in document.tables:
                                for cell in table._cells:
                                    words = cell.text.split()
                                    for word in words:
                                        if word_to_search == word.lower():
                                            count_of_word_in_tables_tmp += 1

                        # To do headers and footers
                        if config.search_headers_and_footers == True:
                            for paragraph in document.sections[0].header.paragraphs:
                                words = paragraph.text.split()
                                for word in words:
                                    if word_to_search == word.lower():
                                        count_of_word_in_headers_and_footers_tmp += 1
                            for paragraph in document.sections[0].footer.paragraphs:
                                words = paragraph.text.split()
                                for word in words:
                                    if word_to_search == word.lower():
                                        count_of_word_in_headers_and_footers_tmp += 1

                        count_of_word_in_document.append(count_of_word_in_text_tmp +
                                                         count_of_word_in_tables_tmp +
                                                         count_of_word_in_headers_and_footers_tmp)

                elif file_type == '.msg':
                    msg = extract_msg.Message(file)
                    sender = msg.sender
                    msg_date = msg.date.replace(',', '')
                    msg_message = msg.body
                    msg_subj = msg.subject

                    # combines message text and subject title
                    msg_text = msg_message + ' ' + msg_subj

                    msg_text = re.sub(r'[^\w\s]', '', msg_text)
                    msg_text = msg_text.replace('\r', ' ')
                    msg_text = msg_text.replace('\n', ' ')
                    words = msg_text.split(' ')
                    words = [word.lower() for word in words]

                    for word_to_search in config.words_to_search:

                        type_of_file.append(file_type)
                        folder_searched.append(os.path.basename(config.path))
                        date_created.append(
                            datetime.strptime(msg_date, '%a %d %b %Y %H:%M:%S %z').strftime('%d/%m/%Y'))
                        date_modified.append('')
                        author.append(sender)
                        document_name.append(file.replace(config.path, ''))
                        word_searched.append(word_to_search)

                        check_dates_match(document_name, date_created, date_modified)
                        count_of_word_in_document_tmp = 0
                        document_name_contains_word.append(document_title_contains_kw(word_to_search, file))

                        for word in words:
                            if word_to_search == word.lower():
                                count_of_word_in_document_tmp += 1

                        count_of_word_in_document.append(count_of_word_in_document_tmp)
            except:
                # if any error happens, add to an error log. Likely to happen when file is open.
                errors.append(f"Error in opening {file.replace(config.path, '')}")
                continue

    # save to dataframe
    df = pd.DataFrame(list(zip(folder_searched, document_name, author, date_created, date_modified,
                               type_of_file, word_searched, document_name_contains_word, count_of_word_in_document)),
                      columns=['folder', 'document_name', 'author', 'date_created', 'date_modified', 'type_of_file',
                               'word_searched', 'document_name_contains_word', 'count_of_word_in_document'])
    print(df)
    df.to_csv(os.path.join(config.path, 'keywords.csv'))

    # save errors to dataframe too
    errors_df = pd.DataFrame({'error': errors})
    print(errors_df)
    errors_df.to_csv(os.path.join(config.path, 'errors.csv'))


if __name__ == '__main__':
    file_count = find_number_of_files()
    keyword_search(file_count)
